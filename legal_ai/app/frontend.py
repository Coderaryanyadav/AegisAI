import streamlit as st
import requests
import json
import os
import time
import pandas as pd
from typing import List, Dict, Any, Optional
from io import BytesIO
try:
    import docx
except ImportError:
    docx = None

# Set page configuration with dark theme default
st.set_page_config(
    page_title="Aegis Legal AI",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# API Server URL
API_URL = "http://127.0.0.1:8000"

# Injected custom CSS for a premium dark SaaS/Glassmorphic interface
def load_css():
    css_path = os.path.join(os.path.dirname(__file__), "styles.css")
    if os.path.exists(css_path):
        with open(css_path, "r") as f:
            css_content = f.read()
    else:
        css_content = ""
        
    # Append custom SaaS-level layout rules, buttons, headers, and hide default streamlit footer
    st.markdown(f"""
    <style>
    {css_content}
    
    /* Hide Streamlit default branding */
    #MainMenu {{visibility: hidden;}}
    footer {{visibility: hidden;}}
    header {{visibility: hidden;}}
    
    /* Document list list item */
    .doc-row {{
        background: #121212;
        border: 1px solid #222222;
        border-radius: 4px;
        padding: 16px 20px;
        margin-bottom: 10px;
        display: flex;
        align-items: center;
        justify-content: space-between;
        transition: all 0.2s ease;
    }}
    .doc-row:hover {{
        background: #181818;
        border-color: #444444;
    }}
    .doc-info {{
        display: flex;
        align-items: center;
        gap: 15px;
    }}
    .doc-icon {{
        font-size: 1.4rem;
        background: #1a1a1a;
        color: #ffffff;
        width: 42px;
        height: 42px;
        border-radius: 2px;
        display: flex;
        align-items: center;
        justify-content: center;
        border: 1px solid #333333;
    }}
    .doc-title {{
        font-weight: 500;
        color: #ffffff;
        font-size: 0.95rem;
    }}
    .doc-meta {{
        font-size: 0.8rem;
        color: #888888;
        margin-top: 2px;
    }}
    
    /* Metric Cards */
    .metric-card {{
        background: #121212;
        border: 1px solid #222222;
        padding: 20px;
        border-radius: 4px;
        text-align: left;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.4);
    }}
    .metric-value {{
        font-size: 2.2rem;
        font-weight: 500;
        color: #ffffff;
        font-family: 'Playfair Display', serif;
        margin-top: 5px;
    }}
    .metric-label {{
        font-size: 0.85rem;
        color: #a3a3a3;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }}
    
    /* Pill Status badges */
    .pill-status {{
        padding: 4px 10px;
        border-radius: 2px;
        font-size: 0.7rem;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }}
    .pill-processed {{
        background: transparent;
        color: #ffffff;
        border: 1px solid #ffffff;
    }}
    .pill-processing {{
        background: transparent;
        color: #888888;
        border: 1px solid #444444;
    }}
    .pill-failed {{
        background: transparent;
        color: #f87171;
        border: 1px solid #ef4444;
    }}
    
    /* Accent text */
    .accent-gradient {{
        color: #ffffff;
        font-weight: 600;
        letter-spacing: 0.02em;
    }}
    </style>
    """, unsafe_allow_html=True)

load_css()

# Session state initialization
if "token" not in st.session_state:
    st.session_state.token = None
if "email" not in st.session_state:
    st.session_state.email = None
if "role" not in st.session_state:
    st.session_state.role = None
if "user_id" not in st.session_state:
    st.session_state.user_id = None
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "drafted_content" not in st.session_state:
    st.session_state.drafted_content = ""

# Helper to verify token and retrieve profile on reload
def load_user_profile():
    if not st.session_state.token:
        return False
    headers = {"Authorization": f"Bearer {st.session_state.token}"}
    try:
        response = requests.get(f"{API_URL}/api/auth/me", headers=headers)
        if response.status_code == 200:
            data = response.json()
            st.session_state.email = data["email"]
            st.session_state.role = data["role"]
            st.session_state.user_id = data["id"]
            return True
        else:
            # Token expired or invalid
            st.session_state.token = None
            st.session_state.email = None
            st.session_state.role = None
            st.session_state.user_id = None
            return False
    except Exception:
        return False

# Authentication Actions
def login_user(email, password):
    payload = {"username": email, "password": password}
    try:
        response = requests.post(f"{API_URL}/api/auth/token", data=payload)
        if response.status_code == 200:
            data = response.json()
            st.session_state.token = data["access_token"]
            st.session_state.role = data["role"]
            st.session_state.email = email
            st.session_state.chat_history = []  # Clear history on fresh login
            st.success("Successfully authenticated.")
            st.rerun()
        else:
            st.error("Invalid credentials.")
    except Exception as e:
        st.error(f"Cannot connect to local auth server: {e}")

def register_user(email, password, role):
    headers = {"Authorization": f"Bearer {st.session_state.token}"}
    params = {"email": email, "password": password, "role": role}
    try:
        response = requests.post(f"{API_URL}/api/auth/register", headers=headers, params=params)
        if response.status_code == 201:
            st.success(f"User account created for {email}.")
        else:
            st.error(response.json().get("detail", "Failed to register user."))
    except Exception as e:
        st.error(f"Error connecting to backend: {e}")

# API Data Actions
def fetch_documents():
    headers = {"Authorization": f"Bearer {st.session_state.token}"}
    try:
        response = requests.get(f"{API_URL}/api/docs", headers=headers)
        if response.status_code == 200:
            return response.json()
    except Exception:
        pass
    return []

def delete_document(doc_id):
    headers = {"Authorization": f"Bearer {st.session_state.token}"}
    try:
        response = requests.delete(f"{API_URL}/api/docs/{doc_id}", headers=headers)
        if response.status_code == 200:
            st.toast("Document permanently removed.")
            return True
        else:
            st.error(response.json().get("detail", "Failed to remove document."))
    except Exception as e:
        st.error(f"Connection error: {e}")
    return False

def upload_document(uploaded_file):
    headers = {"Authorization": f"Bearer {st.session_state.token}"}
    files = {"file": (uploaded_file.name, uploaded_file.getvalue(), uploaded_file.type)}
    try:
        response = requests.post(f"{API_URL}/api/docs/upload", headers=headers, files=files)
        if response.status_code == 202:
            st.toast("Ingestion triggered successfully.", icon="🚀")
            return True
        else:
            st.error(response.json().get("detail", "File upload failed."))
    except Exception as e:
        st.error(f"Upload error: {e}")
    return False

def query_chat(query: str, doc_ids: List[int]):
    headers = {"Authorization": f"Bearer {st.session_state.token}"}
    payload = {"query": query, "document_ids": doc_ids}
    try:
        response = requests.post(f"{API_URL}/api/chat", headers=headers, json=payload)
        if response.status_code == 200:
            return response.json()
        else:
            st.error(response.json().get("detail", "Inference error."))
    except Exception as e:
        st.error(f"Error querying AI engine: {e}")
    return None

def audit_contract(doc_id: int):
    headers = {"Authorization": f"Bearer {st.session_state.token}"}
    try:
        response = requests.post(f"{API_URL}/api/audit-contract/{doc_id}", headers=headers)
        if response.status_code == 200:
            return response.json()
        else:
            st.error(response.json().get("detail", "Risk evaluation crashed."))
    except Exception as e:
        st.error(f"Error running contract audit: {e}")
    return None

def generate_draft(instructions: str, ref_doc_ids: List[int]):
    headers = {"Authorization": f"Bearer {st.session_state.token}"}
    payload = {"instructions": instructions, "reference_doc_ids": ref_doc_ids}
    try:
        response = requests.post(f"{API_URL}/api/draft-document", headers=headers, json=payload)
        if response.status_code == 200:
            return response.json().get("drafted_content", "")
    except Exception as e:
        st.error(f"Drafting error: {e}")
    return ""

def generate_docx_bytes(content: str) -> bytes:
    """Convert text content to a structured .docx file and return raw bytes."""
    bio = BytesIO()
    try:
        if not docx:
            # Fallback if docx module is missing
            st.error("python-docx is not loaded. Cannot export DOCX.")
            return b""
        doc = docx.Document()
        # Add basic styled title
        doc.add_heading("Legal Document Draft", level=0)
        
        # Split text into paragraphs and add them
        paragraphs = content.split("\n")
        for p in paragraphs:
            trimmed = p.strip()
            if not trimmed:
                continue
            if trimmed.startswith("### "):
                doc.add_heading(trimmed.replace("### ", ""), level=2)
            elif trimmed.startswith("## "):
                doc.add_heading(trimmed.replace("## ", ""), level=1)
            elif trimmed.startswith("# "):
                doc.add_heading(trimmed.replace("# ", ""), level=0)
            else:
                doc.add_paragraph(trimmed)
        
        doc.save(bio)
        return bio.getvalue()
    except Exception as e:
        st.error(f"Error generating Word document: {e}")
        return b""

def fetch_audit_logs():
    headers = {"Authorization": f"Bearer {st.session_state.token}"}
    try:
        response = requests.get(f"{API_URL}/api/audit", headers=headers)
        if response.status_code == 200:
            return response.json()
    except Exception as e:
        st.error(f"Failed to fetch audit trail: {e}")
    return []

# ----------------- SYSTEM STATUS DIAGNOSTICS -----------------
def check_system_status() -> Dict[str, Any]:
    status_data = {
        "backend": False,
        "database": "failed",
        "ollama": False,
        "model_pulled": False,
        "model_name": "qwen3:8b"
    }
    # 1. Check Backend & DB
    try:
        r = requests.get(f"{API_URL}/api/health", timeout=1.5)
        if r.status_code == 200:
            res = r.json()
            status_data["backend"] = True
            status_data["database"] = res.get("database", "failed")
            status_data["model_name"] = res.get("llm_model", "qwen3:8b")
    except Exception:
        pass
        
    # 2. Check Ollama & Model presence
    try:
        ollama_url = "http://127.0.0.1:11434"
        r_ollama = requests.get(f"{ollama_url}/api/tags", timeout=1.5)
        if r_ollama.status_code == 200:
            status_data["ollama"] = True
            tags = r_ollama.json().get("models", [])
            model_names = []
            for m in tags:
                name = m.get("name", "")
                model_names.append(name)
                if ":" in name:
                    model_names.append(name.split(":")[0])
            target = status_data["model_name"]
            target_base = target.split(":")[0] if ":" in target else target
            status_data["model_pulled"] = any(target in m or target_base in m for m in model_names)
    except Exception:
        pass
        
    return status_data

def render_diagnostics_sidebar(status_data: Dict[str, Any]):
    st.markdown("<p style='color: #475569; font-size:0.75rem; font-weight:600; text-transform:uppercase; margin-top: 25px; margin-bottom: 8px;'>System Health</p>", unsafe_allow_html=True)
    
    # 1. Backend
    b_color = "🟢" if status_data["backend"] else "🔴"
    b_text = "API Service"
    st.markdown(f"<div style='font-size:0.85rem; display:flex; align-items:center; gap:8px; margin-bottom:4px;'>{b_color} {b_text}</div>", unsafe_allow_html=True)
    
    # 2. Database
    db_color = "🟢" if status_data["database"] == "connected" else "🔴"
    db_text = "Vault Database"
    st.markdown(f"<div style='font-size:0.85rem; display:flex; align-items:center; gap:8px; margin-bottom:4px;'>{db_color} {db_text}</div>", unsafe_allow_html=True)
    
    # 3. Ollama
    o_color = "🟢" if status_data["ollama"] else "🔴"
    o_text = "Ollama Inference"
    st.markdown(f"<div style='font-size:0.85rem; display:flex; align-items:center; gap:8px; margin-bottom:4px;'>{o_color} {o_text}</div>", unsafe_allow_html=True)
    
    # 4. Model
    m_color = "🟢" if status_data["model_pulled"] else "🔴"
    m_text = f"LLM: {status_data['model_name']}"
    st.markdown(f"<div style='font-size:0.85rem; display:flex; align-items:center; gap:8px; margin-bottom:8px;'>{m_color} {m_text}</div>", unsafe_allow_html=True)
    
    # Quick fix message if anything is red
    if not (status_data["backend"] and status_data["ollama"] and status_data["model_pulled"]):
        st.markdown("<div style='margin-top:10px; padding:10px; background:rgba(239, 68, 68, 0.08); border:1px solid rgba(239, 68, 68, 0.2); border-radius:8px; font-size:0.75rem; color:#f87171;'><strong>Diagnostic Alerts:</strong><br/>", unsafe_allow_html=True)
        if not status_data["backend"]:
            st.markdown("- Run start script (./start.sh or start.bat)<br/>", unsafe_allow_html=True)
        if not status_data["ollama"]:
            st.markdown("- Start Ollama desktop application<br/>", unsafe_allow_html=True)
        elif not status_data["model_pulled"]:
            st.markdown(f"- Run: <code>ollama pull {status_data['model_name']}</code> in terminal<br/>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

# Force verify token on reload if present
if st.session_state.token:
    load_user_profile()

status_data = check_system_status()

# ----------------- LOGIN RENDERER -----------------
if not st.session_state.token:
    # Render sidebar with system status even on login page
    with st.sidebar:
        st.markdown("<h2 style='margin-bottom: 2px;'><span class='accent-gradient'>AEGIS</span> AI</h2>", unsafe_allow_html=True)
        st.markdown("<p style='color:#64748b; font-size:0.8rem; margin-top:0;'>Authentication Required</p>", unsafe_allow_html=True)
        render_diagnostics_sidebar(status_data)
        
    st.markdown("<h1 style='text-align: center; margin-top: 80px; font-size: 3rem;'>⚖️ <span class='accent-gradient'>AEGIS</span> LEGAL AI</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #94a3b8; font-size: 1.1rem; margin-bottom: 40px;'>Local, Airtight, and Secure Legal Reasoning Suite</p>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 1.3, 1])
    with col2:
        st.markdown('<div class="legal-card" style="padding: 30px;">', unsafe_allow_html=True)
        st.markdown("<h4 style='margin-top: 0; color: #f8fafc;'>Secure Console Sign-in</h4>", unsafe_allow_html=True)
        st.markdown("<p style='color: #64748b; font-size:0.85rem;'>Verify your firm credentials. All connections are fully local.</p>", unsafe_allow_html=True)
        
        login_email = st.text_input("User Email", key="login_email", placeholder="attorney@firm.local")
        login_password = st.text_input("Security Key / Password", type="password", key="login_password")
        
        st.markdown("<div style='margin-top: 20px;'>", unsafe_allow_html=True)
        if st.button("Authenticate Session", type="primary", use_container_width=True):
            if login_email and login_password:
                login_user(login_email, login_password)
            else:
                st.warning("All authentication fields are required.")
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# ----------------- LOGGED IN RENDERER -----------------

# Define user role badges for side-bar
role_badges = {
    "admin": '<span class="badge-admin" style="font-size:0.75rem; letter-spacing: 0.05em;">👑 ADMIN</span>',
    "lawyer": '<span class="badge-lawyer" style="font-size:0.75rem; letter-spacing: 0.05em;">⚖️ COUNSEL</span>',
    "auditor": '<span class="badge-auditor" style="font-size:0.75rem; letter-spacing: 0.05em;">🔍 AUDITOR</span>'
}

# Fetch documents list globally for convenience
documents = fetch_documents()

# AUTO-REFRESH TRIGGER: If any document is processing, sleep 3s and rerun to update the UI status automatically!
is_processing = any(doc["status"] in ["pending", "processing"] for doc in documents)
if is_processing:
    # Set a tiny notice in side-bar, and trigger rerun
    st.sidebar.warning("⏳ System indexing documents...")
    # Sleep 3 seconds to avoid pegging the database, then trigger rerun
    time.sleep(3)
    st.rerun()

# SIDEBAR NAVIGATION
with st.sidebar:
    st.markdown("<h2 style='margin-bottom: 2px;'><span class='accent-gradient'>AEGIS</span> AI</h2>", unsafe_allow_html=True)
    st.markdown(f"<p style='color:#64748b; font-size:0.8rem; margin-top:0;'>Firm Member: <code>{st.session_state.email}</code></p>", unsafe_allow_html=True)
    st.markdown(f"<div style='margin-bottom: 20px;'>{role_badges.get(st.session_state.role, '')}</div>", unsafe_allow_html=True)
    
    st.markdown("<p style='color: #475569; font-size:0.75rem; font-weight:600; text-transform:uppercase; margin-bottom: 8px;'>Workspace Views</p>", unsafe_allow_html=True)
    
    # Navigation list depending on role
    nav_options = ["📁 Dashboard", "💬 Chat & Q&A", "🔍 Contract Auditor", "✍️ Document Drafting"]
    if st.session_state.role in ["admin", "auditor"]:
        nav_options.append("📋 Security Audit Trail")
    if st.session_state.role == "admin":
        nav_options.append("⚙️ Firm Directory")
        
    choice = st.radio("Navigation Menu", nav_options, label_visibility="collapsed")
    
    # Render Diagnostics in sidebar for authenticated session
    render_diagnostics_sidebar(status_data)
    
    st.markdown("<div style='margin-top: 30px;'>", unsafe_allow_html=True)
    st.markdown("---")
    if st.button("Terminate Session", use_container_width=True, type="secondary"):
        st.session_state.token = None
        st.session_state.email = None
        st.session_state.role = None
        st.session_state.chat_history = []
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

# ----------------- PAGE ROUTING & VIEWS -----------------

# VIEW: DASHBOARD
if choice == "📁 Dashboard":
    st.markdown("<h1>Vault & System Dashboard</h1>", unsafe_allow_html=True)
    st.markdown("<p style='color:#94a3b8;'>Upload and manage securely encrypted legal cases and contracts.</p>", unsafe_allow_html=True)
    
    # Metrics row
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Vault Volume</div>
            <div class="metric-value">{len(documents)}</div>
        </div>
        """, unsafe_allow_html=True)
    with c2:
        indexed_count = sum(1 for d in documents if d["status"] == "processed")
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Vectors Indexed</div>
            <div class="metric-value">{indexed_count}</div>
        </div>
        """, unsafe_allow_html=True)
    with c3:
        processing_count = sum(1 for d in documents if d["status"] in ["pending", "processing"])
        color_class = "color:#fbbf24;" if processing_count > 0 else "color:#34d399;"
        label_text = "Background Jobs" if processing_count > 0 else "System Status"
        value_text = f"{processing_count} Active" if processing_count > 0 else "Online"
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">{label_text}</div>
            <div class="metric-value" style="{color_class}">{value_text}</div>
        </div>
        """, unsafe_allow_html=True)
    with c4:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Vault Security</div>
            <div class="metric-value" style="color:#38bdf8;">AES-256</div>
        </div>
        """, unsafe_allow_html=True)
        
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Layout splits: upload left, file list right
    col_left, col_right = st.columns([1.2, 2])
    
    with col_left:
        st.markdown('<div class="legal-card">', unsafe_allow_html=True)
        st.markdown("<h4 style='margin-top:0; color:#f8fafc;'>Secure Ingest Pipeline</h4>", unsafe_allow_html=True)
        st.markdown("<p style='color:#64748b; font-size:0.85rem;'>Upload documents to be parsed and indexed locally. PDFs are split page-by-page, and text is dynamically encrypted at rest.</p>", unsafe_allow_html=True)
        
        uploaded_file = st.file_uploader("Select PDF or Plain TXT Document:", type=["pdf", "txt"], label_visibility="collapsed")
        
        st.markdown("<div style='margin-top: 15px;'>", unsafe_allow_html=True)
        if st.button("Trigger Vault Ingestion", type="primary", use_container_width=True, disabled=(uploaded_file is None)):
            if uploaded_file is not None:
                if upload_document(uploaded_file):
                    st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
    with col_right:
        col_title, col_sync = st.columns([3, 1])
        with col_title:
            st.markdown("<h3 style='margin-top:0;'>Document Repository</h3>", unsafe_allow_html=True)
        with col_sync:
            if st.button("🔄 Sync Status", use_container_width=True):
                st.rerun()
                
        if not documents:
            st.info("No documents are currently indexed in this workspace.")
        else:
            for doc in documents:
                # Setup status badge
                if doc["status"] == "processed":
                    badge_html = '<span class="pill-status pill-processed">Processed</span>'
                elif doc["status"] in ["pending", "processing"]:
                    badge_html = '<span class="pill-status pill-processing">Processing...</span>'
                else:
                    badge_html = '<span class="pill-status pill-failed">Failed</span>'
                
                # Render clean HTML row
                st.markdown(f"""
                <div class="doc-row">
                    <div class="doc-info">
                        <div class="doc-icon">📄</div>
                        <div>
                            <div class="doc-title">{doc['original_name']}</div>
                            <div class="doc-meta">ID: #{doc['id']} • Uploaded: {doc['uploaded_at'][:16].replace('T', ' ')}</div>
                        </div>
                    </div>
                    <div style="display: flex; align-items: center; gap: 15px;">
                        {badge_html}
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Single-action buttons row below (clean alignment)
                c_del, c_pad = st.columns([1, 6])
                with c_del:
                    # Let admins delete anything, lawyers delete their own
                    if st.session_state.role == "admin" or doc["owner_id"] == st.session_state.user_id:
                        if st.button("Delete Vault File", key=f"del_{doc['id']}", type="secondary", use_container_width=True):
                            if delete_document(doc['id']):
                                st.rerun()
                st.markdown("<div style='margin-bottom:15px;'></div>", unsafe_allow_html=True)

# VIEW: CHAT & Q&A
elif choice == "💬 Chat & Q&A":
    st.markdown("<h1>Semantic Search & Case Assistant</h1>", unsafe_allow_html=True)
    st.markdown("<p style='color:#94a3b8;'>Ask factual questions and receive answers linked to case file citations.</p>", unsafe_allow_html=True)
    
    processed_docs = [doc for doc in documents if doc["status"] == "processed"]
    if not processed_docs:
        st.warning("Please upload and process documents first to start chatting.")
    else:
        # Document context checkboxes in a clean side dashboard
        with st.expander("📂 Scope Search Context", expanded=True):
            st.markdown("<p style='font-size:0.85rem; color:#64748b; margin-top:0;'>Select which case files the AI should analyze (uncheck to exclude):</p>", unsafe_allow_html=True)
            selected_doc_ids = []
            for doc in processed_docs:
                if st.checkbox(doc["original_name"], value=True, key=f"chat_scope_{doc['id']}"):
                    selected_doc_ids.append(doc["id"])
            
        st.markdown("<br>", unsafe_allow_html=True)
        
        # Chat bubbles area
        chat_container = st.container()
        with chat_container:
            for chat in st.session_state.chat_history:
                with st.chat_message(chat["role"]):
                    st.markdown(f"<div style='font-size:0.95rem; color:#f1f5f9;'>{chat['content']}</div>", unsafe_allow_html=True)
                    if chat.get("citations"):
                        with st.expander("📚 Source Document Citations"):
                            for cit in chat["citations"]:
                                st.markdown(f"""
                                <div class="citation-container">
                                    <div class="citation-title">📄 {cit['original_name']} (Page {cit['page_number']})</div>
                                    <div class="citation-text">"...{cit['content']}..."</div>
                                </div>
                                """, unsafe_allow_html=True)
                                
        # Chat input at bottom
        if user_query := st.chat_input("Ask a question about the case documents..."):
            with chat_container:
                # Render query
                with st.chat_message("user"):
                    st.markdown(f"<div style='font-size:0.95rem; color:#f1f5f9;'>{user_query}</div>", unsafe_allow_html=True)
                st.session_state.chat_history.append({"role": "user", "content": user_query})
                
                # Fetch query from AI backend
                with st.chat_message("assistant"):
                    with st.spinner("Analyzing case text chunks..."):
                        response = query_chat(user_query, selected_doc_ids)
                        if response:
                            answer = response["answer"]
                            citations = response["citations"]
                            
                            st.write(answer)
                            if citations:
                                with st.expander("📚 Source Document Citations"):
                                    for cit in citations:
                                        st.markdown(f"""
                                        <div class="citation-container">
                                            <div class="citation-title">📄 {cit['original_name']} (Page {cit['page_number']})</div>
                                            <div class="citation-text">"...{cit['content']}..."</div>
                                        </div>
                                        """, unsafe_allow_html=True)
                            
                            # Append to session state
                            st.session_state.chat_history.append({
                                "role": "assistant",
                                "content": answer,
                                "citations": citations
                            })
                            st.rerun()

# VIEW: CONTRACT AUDITOR
elif choice == "🔍 Contract Auditor":
    st.markdown("<h1>Contract Compliance & Audit Engine</h1>", unsafe_allow_html=True)
    st.markdown("<p style='color:#94a3b8;'>Perform localized clause analysis, severity risk evaluation, and compliance gap detection.</p>", unsafe_allow_html=True)
    
    processed_docs = [doc for doc in documents if doc["status"] == "processed"]
    if not processed_docs:
        st.warning("Please upload contract files (PDF/TXT) to start auditing.")
    else:
        # Select contract to audit in card layout
        st.markdown('<div class="legal-card">', unsafe_allow_html=True)
        doc_options = {doc["id"]: doc["original_name"] for doc in processed_docs}
        selected_doc_id = st.selectbox("Select contract document for risk scan:", options=list(doc_options.keys()), format_func=lambda x: doc_options[x])
        
        st.markdown("<div style='margin-top: 15px;'>", unsafe_allow_html=True)
        trigger_audit = st.button("Perform Automated Clause Scan", type="primary")
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        if trigger_audit:
            with st.spinner("Extracting governing clauses and analyzing risk profiles..."):
                audit_report = audit_contract(selected_doc_id)
                if audit_report:
                    if "error" in audit_report:
                        st.error(audit_report["error"])
                        st.text_area("Raw AI Output Log", audit_report.get("raw_output", ""))
                    else:
                        # Success panel
                        rating = audit_report.get("overall_compliance_rating", "Unknown")
                        rating_color = "#ef4444" if "High" in rating else "#f59e0b" if "Moderate" in rating else "#10b981"
                        
                        st.markdown(f"""
                        <div class="legal-card" style="border-left: 5px solid {rating_color};">
                            <h3 style="margin-top:0; color:#f8fafc;">Risk Analysis Report</h3>
                            <p style="color:#94a3b8; font-size:0.9rem; margin-bottom: 5px;">Overall Compliance Standing:</p>
                            <div style="font-size: 1.5rem; font-weight:700; color:{rating_color};">{rating}</div>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Compute compliance checklist
                        standard_clauses = [
                            {"name": "Governing Law", "key_terms": ["governing", "jurisdiction", "applicable law"]},
                            {"name": "Termination", "key_terms": ["termination", "terminate", "expiration"]},
                            {"name": "Indemnity", "key_terms": ["indemnity", "indemnification", "hold harmless"]},
                            {"name": "Confidentiality", "key_terms": ["confidential", "confidentiality", "non-disclosure"]},
                            {"name": "Force Majeure", "key_terms": ["force majeure", "act of god", "unforeseen event"]},
                            {"name": "Severability", "key_terms": ["severability", "severable", "invalidity"]}
                        ]
                        
                        found_categories = []
                        clauses_list = audit_report.get("clauses_extracted", [])
                        for cl in clauses_list:
                            cat = str(cl.get("clause_type", "")).lower()
                            summary = str(cl.get("summary", "")).lower()
                            found_categories.append(cat)
                            found_categories.append(summary)
                            
                        missing_list = audit_report.get("missing_clauses", [])
                        missing_categories = [str(m.get("clause_type", "")).lower() for m in missing_list]
                        
                        # Generate HTML Checklist
                        checklist_html = '<div class="legal-card" style="padding: 20px;">'
                        checklist_html += '<h4 style="margin-top:0; margin-bottom:15px; color:#ffffff; font-family:\'Playfair Display\', serif;">⚖️ Legal Due Diligence Checklist</h4>'
                        checklist_html += '<div style="display: grid; grid-template-columns: 1fr 1fr; gap: 12px;">'
                        
                        for sc in standard_clauses:
                            sc_name = sc["name"]
                            is_missing = any(sc_name.lower() in mc for mc in missing_categories)
                            is_found = any(sc_name.lower() in fc for fc in found_categories) or any(any(kt in fc for kt in sc["key_terms"]) for fc in found_categories)
                            
                            if is_found and not is_missing:
                                checklist_html += f'<div style="font-size:0.9rem; display:flex; align-items:center; gap:8px; color:#ffffff;">🟢 <strong>{sc_name}</strong>: Identified</div>'
                            else:
                                checklist_html += f'<div style="font-size:0.9rem; display:flex; align-items:center; gap:8px; color:#888888;">🔴 <span style="text-decoration: line-through;">{sc_name}</span>: Omitted</div>'
                                
                        checklist_html += '</div></div>'
                        st.markdown(checklist_html, unsafe_allow_html=True)
                        
                        # Tabs
                        tab_clauses, tab_risks, tab_gaps = st.tabs(["📌 Clauses Extracted", "⚠️ Risk Log", "❌ Compliance Gaps"])
                        
                        with tab_clauses:
                            st.write(f"**Contract Structure:** {audit_report.get('contract_type', 'N/A')}")
                            st.write(f"**Effective Date:** {audit_report.get('effective_date', 'N/A')}")
                            st.write(f"**Identified Signatories:** {', '.join(audit_report.get('parties', []))}")
                            
                            clauses = audit_report.get("clauses_extracted", [])
                            if clauses:
                                pd_clauses = pd.DataFrame(clauses)
                                pd_clauses.columns = ["Clause Category", "Summary of Terms", "Original Text Excerpt"]
                                st.dataframe(pd_clauses, use_container_width=True)
                            else:
                                st.info("No primary clauses extracted.")
                                
                        with tab_risks:
                            risks = audit_report.get("risks_identified", [])
                            if not risks:
                                st.success("No critical liabilities or legal risks identified in scanned text.")
                            else:
                                for r in risks:
                                    severity = r.get("severity", "Low")
                                    color = "#ef4444" if severity == "High" else "#fbbf24" if severity == "Medium" else "#34d399"
                                    bg_color = "rgba(239, 68, 68, 0.08)" if severity == "High" else "rgba(245, 158, 11, 0.08)" if severity == "Medium" else "rgba(16, 185, 129, 0.08)"
                                    
                                    st.markdown(f"""
                                    <div class="legal-card" style="background: {bg_color}; border-color: {color}30; border-left: 4px solid {color}; padding: 18px 24px;">
                                        <div style="display: flex; align-items:center; justify-content:space-between;">
                                            <span style="font-weight:600; color:#f8fafc; font-size:1.05rem;">{r.get('clause_type', 'General Term')}</span>
                                            <span class="pill-status" style="background:{color}20; color:{color}; border: 1px solid {color}40;">{severity} Risk</span>
                                        </div>
                                        <p style="margin-top: 10px; margin-bottom: 8px; color: #cbd5e1; font-size:0.92rem;"><strong>Scanned Liability:</strong> {r.get('description')}</p>
                                        <p style="color: #38bdf8; font-size:0.92rem; margin-top:4px;"><strong>Mitigation Strategy:</strong> {r.get('mitigation')}</p>
                                    </div>
                                    """, unsafe_allow_html=True)
                                    
                        with tab_gaps:
                            gaps = audit_report.get("missing_clauses", [])
                            if not gaps:
                                st.success("No compliance omissions identified.")
                            else:
                                for g in gaps:
                                    st.markdown(f"""
                                    <div class="legal-card" style="border-left: 4px solid #818cf8; padding: 18px 24px;">
                                        <div style="font-weight:600; color:#a5b4fc; font-size:1rem; margin-bottom:6px;">⚠️ Missing clause: {g.get('clause_type')}</div>
                                        <div style="color:#cbd5e1; font-size:0.9rem;">{g.get('explanation')}</div>
                                    </div>
                                    """, unsafe_allow_html=True)

# VIEW: DOCUMENT DRAFTING
elif choice == "✍️ Document Drafting":
    st.markdown("<h1>Legal Draftsman & Builder</h1>", unsafe_allow_html=True)
    st.markdown("<p style='color:#94a3b8;'>Generate formal legal structures, notices, or custom clauses using your files as semantic references.</p>", unsafe_allow_html=True)
    
    col_input, col_output = st.columns([1.1, 1.5])
    
    with col_input:
        st.markdown('<div class="legal-card">', unsafe_allow_html=True)
        st.markdown("<h4 style='margin-top:0; color:#f8fafc;'>Draft Specifications</h4>", unsafe_allow_html=True)
        
        instructions = st.text_area(
            "Drafting Instructions:", 
            placeholder="e.g. 'Draft a tenant eviction notice for non-payment of rent, giving 15 days notice.'",
            height=150
        )
        
        # Scope reference files
        st.markdown("##### Semantic Reference Context")
        if not documents:
            st.info("No documents uploaded. Model will generate a generalized template.")
            ref_doc_ids = []
        else:
            processed_docs = [doc for doc in documents if doc["status"] == "processed"]
            ref_doc_ids = []
            st.markdown("<p style='font-size:0.85rem; color:#64748b; margin-top:0;'>Select reference files to ground this draft (uncheck to exclude):</p>", unsafe_allow_html=True)
            for doc in processed_docs:
                if st.checkbox(doc["original_name"], value=True, key=f"draft_scope_{doc['id']}"):
                    ref_doc_ids.append(doc["id"])
            
        st.markdown("<div style='margin-top: 15px;'>", unsafe_allow_html=True)
        if st.button("Generate Legal Draft", type="primary", use_container_width=True):
            if not instructions:
                st.warning("Please provide drafting instructions first.")
            else:
                with st.spinner("AI drafting document structure..."):
                    drafted_text = generate_draft(instructions, ref_doc_ids)
                    st.session_state.drafted_content = drafted_text
                    st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
    with col_output:
        st.markdown("### Output Viewer")
        if st.session_state.drafted_content:
            st.text_area("Review / Modify Draft:", value=st.session_state.drafted_content, height=450)
            
            # Simple spacing
            st.markdown("<br>", unsafe_allow_html=True)
            col_txt, col_docx = st.columns(2)
            with col_txt:
                st.download_button(
                    label="📥 Download Draft (.txt)",
                    data=st.session_state.drafted_content,
                    file_name="legal_draft_output.txt",
                    mime="text/plain",
                    use_container_width=True
                )
            with col_docx:
                docx_bytes = generate_docx_bytes(st.session_state.drafted_content)
                if docx_bytes:
                    st.download_button(
                        label="📥 Download Draft (.docx)",
                        data=docx_bytes,
                        file_name="legal_draft_output.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
        else:
            st.info("The drafted legal document will appear here after clicking 'Generate Legal Draft'.")

# VIEW: SECURITY AUDIT TRAIL
elif choice == "📋 Security Audit Trail":
    st.markdown("<h1>Security Compliance & Audit Trail</h1>", unsafe_allow_html=True)
    st.markdown("<p style='color:#94a3b8;'>Airtight immutable activity logging tracking firm user actions.</p>", unsafe_allow_html=True)
    
    logs = fetch_audit_logs()
    if not logs:
        st.info("No logs are currently recorded in the audit database.")
    else:
        # Convert to Pandas dataframe
        df_logs = pd.DataFrame(logs)
        df_logs.columns = [
            "Log ID", 
            "User ID", 
            "User Email", 
            "Action Performed", 
            "Target Type", 
            "Target ID", 
            "Timestamp (UTC)", 
            "Metadata / Details", 
            "Origin IP Address"
        ]
        
        # Display clean tabular layout
        st.dataframe(
            df_logs.sort_values(by="Timestamp (UTC)", ascending=False),
            use_container_width=True
        )

# VIEW: FIRM DIRECTORY
elif choice == "⚙️ Firm Directory":
    st.markdown("<h1>Firm Directory & Accounts</h1>", unsafe_allow_html=True)
    st.markdown("<p style='color:#94a3b8;'>Administrative dashboard for employee role management and accounts provisioning.</p>", unsafe_allow_html=True)
    
    col_reg, col_guide = st.columns([1, 1])
    
    with col_reg:
        st.markdown('<div class="legal-card">', unsafe_allow_html=True)
        st.markdown("<h4 style='margin-top:0; color:#f8fafc;'>Provision User Account</h4>", unsafe_allow_html=True)
        reg_email = st.text_input("New Member Email")
        reg_password = st.text_input("Temporary Access Password", type="password")
        reg_role = st.selectbox("Role Permission Level", ["lawyer", "admin", "auditor"])
        
        st.markdown("<div style='margin-top: 15px;'>", unsafe_allow_html=True)
        if st.button("Register Firm Account", type="primary", use_container_width=True):
            if reg_email and reg_password:
                register_user(reg_email, reg_password, reg_role)
            else:
                st.warning("All registration details must be filled out.")
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
    with col_guide:
        st.markdown("### Access Control Schemas")
        st.markdown("""
        - 👑 **Administrator**: Full directory access, vault auditing, and file deletion privileges.
        - ⚖️ **Counsel / Lawyer**: Access to document search, custom legal drafts, contract audits, and case uploads. Can only view/manage files they own.
        - 🔍 **Auditor**: Read-only directory access and security log auditing permissions. Unable to upload or execute model inferences.
        """)
